import sys
from docx import Document
from PyPDF2 import PdfReader
import os
import subprocess
import olefile
import struct
import zlib
import platform


def doc_to_docx(doc_path):
    try:
        system_name = platform.system()
        if system_name == 'Windows':
            libreoffice_path = "C:\\Program Files\\LibreOffice\\program\\soffice.exe"
        elif system_name == 'Linux':
            libreoffice_path = "/usr/bin/libreoffice"
        elif system_name == 'Darwin':
            libreoffice_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        else:
            return None

        command = [
            libreoffice_path,
            "--headless",
            "--convert-to", "docx",
            doc_path,
            "--outdir", os.path.dirname(doc_path)
        ]

        subprocess.run(command, check=True, capture_output=True, text=True)

        docx_file_path = doc_path.replace(".doc", ".docx")

        return docx_file_path
    except Exception as e:
        print(f"doc파일 변환중 오류 발생: {e}", file=sys.stderr)
        return None


def extract_from_docx(doc_or_docx_path):
    """DOCX 파일에서 일반 텍스트와 표 안의 텍스트를 모두 추출"""
    try:
        is_doc = False
        if doc_or_docx_path.endswith('.doc'):
            is_doc = True
            docx_path = doc_to_docx(doc_or_docx_path)
            if doc_or_docx_path is None:
                return "conversion_error"
        else:
            docx_path = doc_or_docx_path

        doc = Document(docx_path)
        text = []

        # 일반 텍스트 추출
        for para in doc.paragraphs:
            if para.text.strip():  # 빈 줄이 아니면 추가
                text.append(para.text)

        # 표 안의 텍스트 추출
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text.append("[표] " + row_text)  # 표 구분자 추가 가능

        if is_doc:
            try:
                if os.path.exists(docx_path):
                    os.remove(docx_path)
                else:
                    print(f"{docx_path} 파일이 존재하지 않습니다.")
            except Exception as e:
                print(f"파일 삭제 중 오류 발생: {e}")

        return "\n".join(text) if text else "no_content"
    except Exception as e:
        print(f"DOCX 파일 처리 중 오류 발생: {e}", file=sys.stderr)
        return "pyerror"


def extract_from_pdf(pdf_path):
    """PDF 파일에서 텍스트 추출"""
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            sub = page.extract_text()
            if sub:
                text += sub + "\n"
        return text.strip()
    except Exception as e:
        print(f"PDF 파일 처리 중 오류 발생: {e}", file=sys.stderr)
        return "pyerror"


def extract_from_hwp(hwp_path):
    try:
        f = olefile.OleFileIO(hwp_path)
        dirs = f.listdir()

        # HWP 파일 검증
        if ["FileHeader"] not in dirs or ["\x05HwpSummaryInformation"] not in dirs:
            raise Exception("Not Valid HWP.")

        # 문서 포맷 압축 여부 확인
        header = f.openstream("FileHeader")
        header_data = header.read()
        is_compressed = (header_data[36] & 1) == 1

        # Body Sections 불러오기
        nums = []
        for d in dirs:
            if d[0] == "BodyText":
                nums.append(int(d[1][len("Section"):]))
        sections = ["BodyText/Section" + str(x) for x in sorted(nums)]

        # 전체 text추출
        text = ""
        for section in sections:
            body_text = f.openstream(section)
            data = body_text.read()
            if is_compressed:
                unpack_data = zlib.decompress(data, -15)
            else:
                unpack_data = data

            # 각 section내 text추출
            section_text = ""
            i = 0
            size = len(unpack_data)
            while i < size:
                header = struct.unpack_from("<I", unpack_data, i)[0]
                rec_type = header & 0x3ff
                rec_len = (header >> 20) & 0xfff

                if rec_type in [67]:
                    rec_data = unpack_data[i + 4:i + 4 + rec_len]
                    section_text += rec_data.decode("utf-16")
                    section_text += "\n"

                i += 4 + rec_len

            text += section_text
            text += "\n"

        return text
    except Exception as e:
        print(f"HWP 파일 처리 중 오류 발생: {e}", file=sys.stderr)
        return "pyerror"


def extract(file_type, file_path):
    """파일 유형에 따라 적절한 텍스트 추출 함수 호출"""
    if file_type == "docx" or file_type == "doc":
        return extract_from_docx(file_path)
    elif file_type == "pdf":
        return extract_from_pdf(file_path)
    elif file_type == "hwp":
        return extract_from_hwp(file_path)
    else:
        print(f"지원하지 않는 파일 유형입니다: {file_type}", file=sys.stderr)
        return "pyerror"


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("사용법: python extract_text.py [파일유형] [파일경로]", file=sys.stderr)
        sys.exit(1)

    file_type = sys.argv[1]  # 파일 확장자 (예: pdf, docx, hwp)
    file_path = sys.argv[2]

    # file_type = "doc"
    # file_path = "이력서 및 자기소개서.doc"

    content = extract(file_type, file_path)
    print(content)
    # print(content.encode('utf-8').decode('utf-8'))  # UTF-8 인코딩 강제
